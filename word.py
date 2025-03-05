import os
from docx import Document
import xml.etree.ElementTree as ET

def escapar_caracteres(texto):
    """Escapa caracteres que no son compatibles con XML."""
    try:
        ET.fromstring(f'<raiz>{texto}</raiz>')
        return texto
    except ET.ParseError:
        # Elimina o reemplaza caracteres no válidos
        return ''.join(char for char in texto if ord(char) > 31 or char in '\t\n\r')

def agregar_archivo_a_word(ruta_directorio, nombre_documento_salida):
    doc = Document()
    
    archivos_encontrados = 0
    # Ajusta las extensiones según los tipos de archivos que esperas en src
    extensiones = ('.js', '.jsx', '.css', '.json', '.html')

    # Recorre solo la carpeta 'src'
    for root, dirs, files in os.walk(ruta_directorio):
        for filename in files:
            if filename.endswith(extensiones):
                archivos_encontrados += 1
                filepath = os.path.join(root, filename)
                with open(filepath, 'r', encoding='utf-8') as file:
                    code = file.read()

                # Limpia el código antes de añadirlo al documento
                code = escapar_caracteres(code)

                doc.add_paragraph(f"Ruta del archivo: {filepath}")
                doc.add_paragraph("Contenido del archivo:")
                doc.add_paragraph(code)
                doc.add_page_break()

    print(f"Archivos procesados: {archivos_encontrados}")
    if archivos_encontrados > 0:
        doc.save(nombre_documento_salida)
        print(f"Documento guardado: {nombre_documento_salida}")
    else:
        print("No se encontraron archivos para procesar.")

# Define la ruta directamente a la carpeta 'src'
directorio = '/Users/luisrojas/Documents/GitHub/Protfolio2/src'
nombre_doc_salida = 'frontreact.docx'
agregar_archivo_a_word(directorio, nombre_doc_salida)
